from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUTPUT_DOCX = DOCS_DIR / "程序鉴别材料_软著版.docx"
OUTPUT_PDF = DOCS_DIR / "程序鉴别材料_软著版.pdf"
OUTPUT_MANIFEST = DOCS_DIR / "程序鉴别材料_软著版_校验清单.txt"

LINES_PER_PAGE = 50
TOTAL_PAGES = 60
HEADER_TEXT = "LLMGather 大模型智能体聚合平台 V1.0"
PDF_FONT_NAME = "EmbeddedCJK"
PDF_FONT_PATH = "/System/Library/Fonts/Supplemental/Songti.ttc"
PDF_FONT_SIZE = 7.4
HEADER_FONT_SIZE = 9
FOOTER_FONT_SIZE = 8
PAGE_WIDTH, PAGE_HEIGHT = A4
LEFT_MARGIN = 36
RIGHT_MARGIN = 36
TOP_MARGIN = 46
HEADER_Y = PAGE_HEIGHT - 24
FOOTER_Y = 18
LINE_HEIGHT = 14.85
MAX_TEXT_WIDTH = PAGE_WIDTH - LEFT_MARGIN - RIGHT_MARGIN
SAFE_TEXT_WIDTH = MAX_TEXT_WIDTH - 20
DISPLAY_PREFIX_SAMPLE = "50 "

pdfmetrics.registerFont(TTFont(PDF_FONT_NAME, PDF_FONT_PATH))

SEGMENTS = [
    [
        "frontend/src/components/app/pages/MultimodalModule.vue",
        "frontend/src/components/app/pages/ConsoleModule.vue",
        "frontend/src/components/app/pages/CollabModule.vue",
        "frontend/src/types/agent.ts",
        "frontend/src/components/app/pages/GroupModule.vue",
    ],
    [
        "backend/src/agents/agents.service.spec.ts",
        "backend/scripts/export-sanitized-sql.cjs",
        "backend/src/providers/gemini.provider.spec.ts",
        "backend/src/memory/providers/langgraph-memory.provider.ts",
    ],
]


def pdf_text_width(text: str) -> float:
    return pdfmetrics.stringWidth(text, PDF_FONT_NAME, PDF_FONT_SIZE)


def break_index_by_pdf_width(text: str, max_width: float) -> int:
    if pdf_text_width(text) <= max_width:
        return len(text)
    low = 1
    high = len(text)
    best = 1
    while low <= high:
        mid = (low + high) // 2
        if pdf_text_width(text[:mid]) <= max_width:
            best = mid
            low = mid + 1
        else:
            high = mid - 1
    break_marks = [" ", ",", ";", ")", "}", "]", ">", "|", ":", ".", "/", "-", "_"]
    marked = max(text.rfind(mark, 0, best + 1) for mark in break_marks)
    if marked >= max(12, best // 2):
        candidate = marked + 1
        if pdf_text_width(text[:candidate]) <= max_width:
            return candidate
    return best


def wrap_code_line(text: str, max_width: float) -> list[str]:
    if text == "":
        return [""]
    chunks: list[str] = []
    current = text
    while pdf_text_width(current) > max_width:
        cut = break_index_by_pdf_width(current, max_width)
        chunks.append(current[:cut].rstrip())
        current = "    " + current[cut:].lstrip()
    chunks.append(current)
    return chunks


def material_lines_for_file(relative_path: str) -> list[str]:
    path = ROOT / relative_path
    raw_lines = path.read_text(encoding="utf-8").splitlines()
    lines = [f"// ====== {relative_path} ======"]
    display_prefix_width = pdf_text_width(DISPLAY_PREFIX_SAMPLE)
    for raw in raw_lines:
        normalized = raw.replace("\t", "  ")
        wrapped = wrap_code_line(normalized, SAFE_TEXT_WIDTH - display_prefix_width)
        lines.append(wrapped[0])
        for continuation in wrapped[1:]:
            lines.append(continuation)
    return lines


def build_material() -> tuple[list[str], list[dict[str, object]]]:
    all_lines: list[str] = []
    manifest: list[dict[str, object]] = []
    cursor = 0
    for segment_index, segment_files in enumerate(SEGMENTS, start=1):
        for relative_path in segment_files:
            file_lines = material_lines_for_file(relative_path)
            file_start = cursor + 1
            all_lines.extend(file_lines)
            cursor += len(file_lines)
            manifest.append(
                {
                    "segment": segment_index,
                    "path": relative_path,
                    "line_count": len(file_lines),
                    "start": file_start,
                    "end": cursor,
                    "start_page": ((file_start - 1) // LINES_PER_PAGE) + 1,
                    "end_page": ((cursor - 1) // LINES_PER_PAGE) + 1,
                }
            )
    expected_total = TOTAL_PAGES * LINES_PER_PAGE
    if len(all_lines) != expected_total:
        raise RuntimeError(f"总行数错误：{len(all_lines)}，期望 {expected_total}。")
    if not all_lines[-1].strip():
        raise RuntimeError("第60页末尾是空行，不符合完整程序结束要求。")
    return all_lines, manifest


def set_docx_margins(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.0)


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def strip_docx_body_spacing(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(7)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(10.8)


def add_no_wrap(paragraph) -> None:
    paragraph_property = paragraph._p.get_or_add_pPr()
    word_wrap = OxmlElement("w:wordWrap")
    word_wrap.set(qn("w:val"), "0")
    paragraph_property.append(word_wrap)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run("第 ")
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instr)
    run._r.append(field_end)
    paragraph.add_run(f" 页 / 共 {TOTAL_PAGES} 页")


def setup_docx_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_paragraph = section.header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_paragraph.add_run(HEADER_TEXT)
    set_east_asia_font(header_run, "宋体")
    header_run.font.size = Pt(9)

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(footer_paragraph)
    for run in footer_paragraph.runs:
        set_east_asia_font(run, "宋体")
        run.font.size = Pt(8)


def display_line(global_index: int, line: str) -> str:
    page_line_number = ((global_index - 1) % LINES_PER_PAGE) + 1
    return f"{page_line_number:>2} {line}"


def write_docx(lines: list[str]) -> None:
    doc = Document()
    set_docx_margins(doc)
    setup_docx_header_footer(doc)
    strip_docx_body_spacing(doc)
    for index, line in enumerate(lines, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(10.8)
        add_no_wrap(paragraph)
        run = paragraph.add_run(display_line(index, line))
        set_east_asia_font(run, "宋体")
        run.font.size = Pt(7)
        if index % LINES_PER_PAGE == 0 and index != len(lines):
            run.add_break(WD_BREAK.PAGE)
    doc.save(OUTPUT_DOCX)


def write_pdf(lines: list[str]) -> None:
    pdf = canvas.Canvas(str(OUTPUT_PDF), pagesize=A4)
    pdf.setTitle("程序鉴别材料_软著版")
    for page_index in range(TOTAL_PAGES):
        page_number = page_index + 1
        pdf.setFont(PDF_FONT_NAME, HEADER_FONT_SIZE)
        pdf.drawCentredString(PAGE_WIDTH / 2, HEADER_Y, HEADER_TEXT)
        pdf.setFont(PDF_FONT_NAME, FOOTER_FONT_SIZE)
        pdf.drawCentredString(PAGE_WIDTH / 2, FOOTER_Y, f"第 {page_number} 页 / 共 {TOTAL_PAGES} 页")
        pdf.setFont(PDF_FONT_NAME, PDF_FONT_SIZE)
        start = page_index * LINES_PER_PAGE
        page_lines = lines[start : start + LINES_PER_PAGE]
        if len(page_lines) != LINES_PER_PAGE:
            raise RuntimeError(f"第 {page_number} 页行数不是 {LINES_PER_PAGE}。")
        for row_index, line in enumerate(page_lines):
            rendered_line = display_line(start + row_index + 1, line)
            if pdf_text_width(rendered_line) > SAFE_TEXT_WIDTH + 0.01:
                raise RuntimeError(f"第 {page_number} 页第 {row_index + 1} 行超出页面宽度。")
            y = PAGE_HEIGHT - TOP_MARGIN - (row_index + 1) * LINE_HEIGHT
            pdf.drawString(LEFT_MARGIN, y, rendered_line)
        pdf.showPage()
    pdf.save()


def write_manifest(lines: list[str], manifest: list[dict[str, object]]) -> None:
    page_30_tail = display_line(30 * LINES_PER_PAGE, lines[30 * LINES_PER_PAGE - 1])
    page_60_tail = display_line(60 * LINES_PER_PAGE, lines[60 * LINES_PER_PAGE - 1])
    rows = [
        "程序鉴别材料软著版校验清单",
        f"总页数：{TOTAL_PAGES}",
        f"每页行数：{LINES_PER_PAGE}",
        f"页眉：{HEADER_TEXT}",
        "页码：居中页脚",
        "行号：每页独立从 1 到 50",
        f"总代码行：{len(lines)}",
        f"第30页末尾：{page_30_tail}",
        f"第60页末尾：{page_60_tail}",
        "",
        "文件清单：",
    ]
    for item in manifest:
        rows.append(
            f"第{item['segment']}段 | {item['path']} | "
            f"{item['line_count']}行 | 全局行 {item['start']}-{item['end']} | "
            f"页 {item['start_page']}-{item['end_page']}"
        )
    rows.append("")
    rows.append("校验结论：第60页末尾为完整源码文件末尾，没有用空白行补齐；第30页允许连续分页。")
    OUTPUT_MANIFEST.write_text("\n".join(rows) + "\n", encoding="utf-8")


def main() -> None:
    lines, manifest = build_material()
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    write_docx(lines)
    write_pdf(lines)
    write_manifest(lines, manifest)
    print(f"wrote {OUTPUT_DOCX}")
    print(f"wrote {OUTPUT_PDF}")
    print(f"wrote {OUTPUT_MANIFEST}")


if __name__ == "__main__":
    main()
